from docx import Document
from docx.shared import Inches

document = Document('document.docx')

document.add_picture('demo.png', width=Inches(1.25))
paragraph = document.add_paragraph()
paragraph.add_run('Текстовая подпись под изображением')
document.save('document.docx')
