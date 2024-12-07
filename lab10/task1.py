from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

document.add_paragraph('В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти: ')

paragraph = document.add_paragraph(
    'Флеш-память: используется для хранения скетчей', style='List Bullet'
)
paragraph.paragraph_format.left_indent = Inches(1)
paragraph.paragraph_format.space_after = Pt(0)

paragraph = document.add_paragraph(
    'ОЗУ (', style='List Bullet'
)
paragraph.add_run('SRAM').bold = True
paragraph.add_run('-')
paragraph.add_run('static random access memory').italic = True
paragraph.add_run(', статическая оперативная память с произвольным доступом): используется для хранения и работы переменных.')
paragraph.paragraph_format.left_indent = Inches(1)
paragraph.paragraph_format.space_before = Pt(0)
paragraph.paragraph_format.space_after = Pt(0)

paragraph = document.add_paragraph(
    'EEPROM (энергонезависимая память): используется для хранения постоянной информации.', style='List Bullet'
)
paragraph.paragraph_format.left_indent = Inches(1)

document.add_paragraph('Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.')

records = (
    ('16 Кбайт', '32 Кбайт', '128 Кбайт', '256 Кбайт'),
    ('1 Кбайт', '2 Кбайт', '8 Кбайт', '8 Кбайт'),
    ('512 байт', '1024 байта', '4 Кбайта', '4 Кбайта')
)

columns = ('ATmega168', 'ATmega328', 'ATmega1280', 'ATmega2560')
rows = ('Flash (1 кБ flash-памяти занят загрузчиком)', 'SRAM', 'EEPROM')

table = document.add_table(rows=1, cols=5)

hdr_cells = table.rows[0].cells
for i, column in enumerate(columns):
    hdr_cells[i + 1].text = column
    hdr_cells[i + 1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, record in enumerate(records):
    row_cells = table.add_row().cells
    row_cells[0].text = rows[i] 
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, value in enumerate(record):
        row_cells[j + 1].text = value
        row_cells[j + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def fillCellWithGreyColor(cell_xml_element):
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), "eeeeee")
    table_cell_properties.append(shade_obj)


for i in range(4):
    fillCellWithGreyColor(table.rows[i].cells[0]._tc)
    fillCellWithGreyColor(table.rows[0].cells[i + 1]._tc)
    

paragraph = document.add_paragraph()
paragraph.add_run('Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25 С. Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM.').italic = True
paragraph.paragraph_format.space_before = Pt(12)

document.save('document.docx')